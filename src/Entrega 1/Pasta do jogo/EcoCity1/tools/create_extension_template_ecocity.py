from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


SOURCE = Path(r"C:\Users\admin\Downloads\Documento - Projeto de Extensão - COM Empresa - 2026_1 (1).docx")
OUTPUT = Path(r"C:\Users\admin\Downloads\Documento - Projeto de Extensão - COM Empresa - Eco City.docx")


def set_font_run(run, size=None, bold=None, color=None):
    run.font.name = "Arial"
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def set_cell_text(cell, text, bold_first=False):
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)

    if not cell.paragraphs:
        cell.add_paragraph()

    clear_paragraph(cell.paragraphs[0])
    lines = text.split("\n")

    for idx, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        if idx > 0:
            clear_paragraph(paragraph)
        run = paragraph.add_run(line)
        set_font_run(run, size=10.5, bold=(bold_first and idx == 0), color=(0, 0, 0))


def force_document_text_dark(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_font_run(run, color=(0, 0, 0))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_font_run(run, color=(0, 0, 0))


def main():
    doc = Document(SOURCE)
    tables = doc.tables

    set_cell_text(tables[0].cell(0, 0), "Eco City: jogo digital para conscientização sobre gestão urbana sustentável e tomada de decisão estratégica.")

    team_rows = [
        ("Arthur Lima De Luiz", "26028799"),
        ("Brian Walter", "26029357"),
        ("Leonardo Batista Franca", "26028512"),
    ]
    for idx, (name, ra) in enumerate(team_rows, start=1):
        set_cell_text(tables[1].cell(idx, 0), name)
        set_cell_text(tables[1].cell(idx, 1), ra)
    for idx in range(len(team_rows) + 1, len(tables[1].rows)):
        set_cell_text(tables[1].cell(idx, 0), "")
        set_cell_text(tables[1].cell(idx, 1), "")

    set_cell_text(
        tables[2].cell(0, 0),
        "Professores orientadores: Victor Bruno Alexander Rosetti de Quiroz, Renata Muniz Do Nascimento, Adriano Felix Valente, Eduardo Savino Gomes e Luis Fernando dos Santos Pires.",
    )
    set_cell_text(tables[3].cell(0, 0), "Ciência da Computação.")
    set_cell_text(tables[4].cell(0, 0), "Projeto Interdisciplinar:")
    set_cell_text(
        tables[4].cell(0, 1),
        "Aplicação de tecnologia educacional com foco em sustentabilidade urbana, planejamento estratégico e conscientização cidadã por meio de um jogo digital.",
    )
    set_cell_text(
        tables[5].cell(0, 0),
        "Atividade de Extensão não implementada na prática (proposta de intervenção).",
    )
    set_cell_text(
        tables[6].cell(0, 0),
        "ODS 11 - Cidades e Comunidades Sustentáveis; ODS 4 - Educação de Qualidade; ODS 9 - Indústria, Inovação e Infraestrutura; ODS 13 - Ação Contra a Mudança Global do Clima.",
    )
    set_cell_text(
        tables[7].cell(0, 0),
        "Produto extensionista: jogo digital 3D Eco City, desenvolvido na Unity, com tabuleiro estratégico, dado físico com carregamento, movimentação de peão, sistema de compra de propriedades, eventos urbanos, construções visuais, tela final classificatória e cenário urbano de fundo. Evidências previstas: capturas de tela, build executável, código-fonte, GDD e documentação do projeto. Assets utilizados: Board Game Essentials (dado e peão), Colorful City (skyline e vegetação), Low Poly Houses Free Pack (construções e detalhes urbanos), Simple Building Generic Free (escola e hospital), SolarPanel (usina solar) e SamhuiRestaurant (praça de alimentação).",
    )
    set_cell_text(
        tables[8].cell(0, 0),
        "A intervenção pode ser aplicada em escolas de ensino médio, cursos técnicos, feiras acadêmicas, semanas de extensão e oficinas introdutórias sobre sustentabilidade e cidadania. O jogo foi pensado para um contexto educacional e demonstrativo, com execução viável em laboratórios de informática, mostras universitárias e apresentações com projetor.",
    )
    set_cell_text(
        tables[9].cell(0, 0),
        "O público-alvo é composto por estudantes de 15 a 25 anos, comunidade escolar, visitantes de eventos acadêmicos e pessoas interessadas em gestão urbana e sustentabilidade. Trata-se de um grupo com diferentes níveis de familiaridade com planejamento urbano, o que reforça a importância de uma abordagem lúdica, visual e acessível.",
    )
    set_cell_text(
        tables[10].cell(0, 0),
        "O problema central identificado é a dificuldade de compreender, de forma prática e acessível, como decisões urbanas afetam simultaneamente economia, qualidade de vida e meio ambiente. Em geral, esses temas são discutidos de forma abstrata, sem ferramentas de experimentação. O projeto responde a esse problema por meio de um jogo que transforma esses conceitos em escolhas concretas e consequências imediatas.",
    )
    set_cell_text(
        tables[11].cell(0, 0),
        "Hipótese de intervenção: um jogo digital educativo, com regras simples e feedback visual claro, pode ampliar o entendimento do público sobre gestão urbana sustentável. Ao experimentar decisões de compra, renda, poluição, bem-estar e eventos da cidade, o participante tende a perceber melhor os impactos de políticas urbanas e a importância do equilíbrio entre desenvolvimento econômico e responsabilidade socioambiental.",
    )
    set_cell_text(
        tables[12].cell(0, 0),
        "Eco City é uma proposta de projeto de extensão baseada no desenvolvimento de um jogo digital 3D de estratégia, criado na Unity, para sensibilizar estudantes e a comunidade sobre sustentabilidade urbana. O jogo coloca o participante no papel de gestor de uma cidade, que precisa tomar decisões ao longo de 12 rodadas em um tabuleiro com 20 casas. Cada propriedade adquirida gera efeitos em dinheiro, bem-estar e poluição, além de eventos especiais e múltiplos finais. A proposta pode ser aplicada em oficinas, mostras acadêmicas e ações educativas, funcionando como ferramenta lúdica de discussão sobre planejamento urbano, equilíbrio ambiental e responsabilidade pública.",
    )
    set_cell_text(
        tables[13].cell(0, 0),
        "A extensão universitária busca aproximar o conhecimento acadêmico das demandas sociais, criando intervenções com impacto formativo e comunitário. Nesse contexto, Eco City propõe um jogo digital como instrumento educativo voltado à compreensão de desafios urbanos contemporâneos. O projeto dialoga diretamente com a ODS 11, ao abordar cidades sustentáveis, e também com a ODS 4, ao favorecer aprendizagem significativa, com a ODS 9, por estimular inovação tecnológica, e com a ODS 13, ao discutir impactos ambientais. A fundamentação do projeto parte da ideia de que jogos sérios e experiências interativas podem ampliar o engajamento dos participantes, facilitando a compreensão de sistemas complexos e relações de causa e efeito. Em Eco City, o jogador observa na prática como escolhas econômicas, sociais e ambientais precisam ser equilibradas para produzir uma cidade saudável e funcional.",
    )
    set_cell_text(
        tables[14].cell(0, 0),
        "• Desenvolver um jogo digital 3D com foco em gestão urbana sustentável.\n• Simular decisões sobre propriedades, renda, poluição e bem-estar em um tabuleiro estratégico.\n• Estimular o pensamento crítico sobre políticas públicas urbanas e sustentabilidade.\n• Produzir um artefato extensionista aplicável em contextos educativos e acadêmicos.\n• Promover discussão sobre equilíbrio entre desenvolvimento econômico, qualidade de vida e preservação ambiental.",
    )
    set_cell_text(
        tables[15].cell(0, 0),
        "A metodologia combina pesquisa, design de jogo, implementação técnica e validação prática. Inicialmente, a equipe levantou referências sobre sustentabilidade urbana, jogos educativos e mecânicas de tabuleiro. Em seguida, estruturou o fluxo principal do jogo, definindo atributos, rodada, progressão, eventos especiais e finais classificatórios. A implementação foi realizada na Unity com C#, incluindo dado físico com carregamento, movimentação do peão, compra de propriedades, construções visuais, renda recorrente, HUD, menus, câmera dinâmica e tela final. Na composição visual, foram integrados assets de terceiros para apoiar o protótipo, como peão e dado, prédios low poly, árvores, escola, hospital, painel solar e restaurante temático. Para a etapa extensionista, a proposta é utilizar o jogo em demonstrações, oficinas e apresentações, permitindo que os participantes joguem, discutam as consequências das decisões e relacionem a experiência com problemas reais das cidades. A coleta de percepção pode ocorrer por observação, conversa guiada e questionários simples após a atividade.",
    )
    set_cell_text(
        tables[16].cell(0, 0),
        "Resultados esperados: maior compreensão do público sobre os efeitos interligados de decisões urbanas; fortalecimento do debate sobre sustentabilidade, cidadania e planejamento urbano; produção de um jogo educacional funcional e apresentável; e geração de material acadêmico que comprove o desenvolvimento do projeto. Como resultado parcial já alcançado, o jogo possui menu inicial, HUD, dado físico, movimentação do peão, compra de propriedades, eventos de cidade, cenário urbano, sistema de pontuação e tela final classificatória.",
    )
    set_cell_text(
        tables[17].cell(0, 0),
        "Conclui-se que Eco City constitui uma proposta consistente de extensão universitária por transformar conceitos de sustentabilidade urbana em uma experiência prática e interativa. O projeto responde ao problema central ao criar uma ferramenta acessível para discutir planejamento urbano, impacto ambiental e tomada de decisão estratégica. Além do valor técnico no desenvolvimento do software, o jogo apresenta potencial educativo e extensionista para futuras ações em escolas, feiras e eventos acadêmicos.",
    )
    set_cell_text(
        tables[18].cell(0, 0),
        "NAÇÕES UNIDAS. Objetivos de Desenvolvimento Sustentável. Disponível em: https://brasil.un.org/pt-br/sdgs.\nUNITY TECHNOLOGIES. Unity Manual e Scripting API. Disponível em: https://docs.unity3d.com/.\nSALEN, Katie; ZIMMERMAN, Eric. Rules of Play: Game Design Fundamentals.\nReferências complementares sobre jogos educativos, sustentabilidade urbana e aprendizagem baseada em simulação.",
    )
    set_cell_text(
        tables[19].cell(0, 0),
        "Produto extensionista principal: game digital educativo Eco City.\nProdutos complementares: GDD atualizado, documentação do projeto, imagens de tela, possível build executável, apresentação acadêmica e roteiro para demonstração em oficinas e mostras.",
    )
    set_cell_text(tables[20].cell(0, 0), "Fontes:\n• Objetivos de Desenvolvimento Sustentável\n• Documentação oficial da Unity\n• Referências de game design e jogos sérios")
    set_cell_text(tables[20].cell(0, 1), "Links:\nhttps://brasil.un.org/pt-br/sdgs\nhttps://docs.unity3d.com/\nhttps://unity.com/")
    set_cell_text(tables[21].cell(0, 1), "https://www.fecap.br/")
    set_cell_text(tables[21].cell(1, 1), "https://www.fecap.br/")

    force_document_text_dark(doc)
    doc.save(OUTPUT)
    print(str(OUTPUT))


if __name__ == "__main__":
    main()
