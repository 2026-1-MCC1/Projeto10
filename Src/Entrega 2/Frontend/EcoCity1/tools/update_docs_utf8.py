from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


ROOT = Path(r"C:\Users\admin\Downloads\projetos\Projeto10\src\Entrega PI\Frontend\EcoCity1")
GDD_SOURCE_PATH = next(ROOT.glob("GDD EcoCity*.docx"))
EXT_SOURCE_PATH = next(ROOT.glob("Projeto de Extensão*.docx"))
GDD_PATH = ROOT / "GDD EcoCity.docx"
EXT_PATH = ROOT / "Projeto de Extensão Eco City - COM Empresa.docx"
GDD_OUT_PATH = ROOT / "GDD EcoCity - Atualizado.docx"
EXT_OUT_PATH = ROOT / "Projeto de Extensão Eco City - COM Empresa - Atualizado.docx"


def set_font_run(run, size=None, bold=None):
    run.font.name = "Arial"
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def set_paragraph_text(paragraph, text, bold=False, size=12, align=None):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_font_run(run, size=size, bold=bold)
    if align is not None:
        paragraph.alignment = align


def set_cell_text(cell, text, bold_first=False):
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)

    if not cell.paragraphs:
        cell.add_paragraph()

    clear_paragraph(cell.paragraphs[0])
    lines = text.split("\n")

    for idx, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        if idx > 0:
            clear_paragraph(paragraph)
        if line.startswith("• "):
            try:
                paragraph.style = "List Bullet"
            except Exception:
                pass
            line = line[2:]
        run = paragraph.add_run(line)
        set_font_run(run, size=10.5, bold=(bold_first and idx == 0))


def update_gdd():
    doc = Document(GDD_SOURCE_PATH)
    updates = {
        7: ("ECO CITY", True, 18, WD_PARAGRAPH_ALIGNMENT.CENTER),
        8: ("DOCUMENTO DE DESIGN DE GAME (GDD)", True, 14, WD_PARAGRAPH_ALIGNMENT.CENTER),
        19: ("Arthur Lima De Luiz\nBrian Walter\nLeonardo Batista Franca", False, 12, None),
        32: ("Eco City é um jogo digital 3D, single player, desenvolvido na Unity com C#, que simula a gestão de uma cidade sustentável por meio de um tabuleiro estratégico. Ao longo de 12 rodadas, o jogador percorre 20 casas, compra propriedades, administra dinheiro, bem-estar e poluição, enfrenta eventos urbanos e recebe uma classificação final conforme o equilíbrio obtido na cidade.", False, 11, None),
        34: ("Proporcionar uma experiência estratégica baseada em escolhas com consequências econômicas, sociais e ambientais.", False, 11, None),
        35: ("Estimular a compreensão de gestão urbana sustentável por meio de mecânicas acessíveis e visuais.", False, 11, None),
        36: ("Incentivar pensamento estratégico, análise de trade-offs e planejamento de longo prazo.", False, 11, None),
        37: ("Aplicar conceitos de programação, modelagem matemática, física, UI e game design em um produto jogável.", False, 11, None),
        39: ("Jogo 3D desenvolvido na Unity.", False, 11, None),
        40: ("Sistema baseado em turnos com 12 rodadas.", False, 11, None),
        41: ("Tabuleiro 6x6 com 20 casas jogáveis ao redor da borda.", False, 11, None),
        42: ("Sistema de compra opcional de propriedades com impactos distintos.", False, 11, None),
        44: ("Jovens e adultos a partir de 15 anos, especialmente estudantes interessados em estratégia, sustentabilidade e gestão pública.", False, 11, None),
        46: ("Sistema de decisão estratégica baseado em equilíbrio entre dinheiro, bem-estar e poluição.", False, 11, None),
        48: ("Forças: tema atual, escopo completo e feedback visual claro.", False, 11, None),
        49: ("Fraquezas: experiência single player e necessidade de calibragem fina de balanceamento.", False, 11, None),
        50: ("Oportunidades: uso educacional, demonstrações acadêmicas e expansão futura com mais eventos e multiplayer.", False, 11, None),
        51: ("Ameaças: tempo de produção limitado e dependência de polimento final para apresentação.", False, 11, None),
        53: ("Eco City oferece uma experiência interativa que transforma conceitos de sustentabilidade urbana em decisões concretas de jogo, aproximando o jogador da realidade de um gestor municipal de forma acessível, visual e estratégica.", False, 11, None),
        55: ("Sistema de turnos com progressão por rodadas.", False, 11, None),
        56: ("Sistema econômico com compra, aluguel e renda recorrente.", False, 11, None),
        57: ("Sistema de atributos: dinheiro, bem-estar, poluição e pontuação.", False, 11, None),
        58: ("Dado físico com carregamento de força e lançamento pela câmera.", False, 11, None),
        59: ("Movimentação do peão com salto fluido entre tiles.", False, 11, None),
        64: ("O jogador assume o papel de um gestor urbano responsável por desenvolver uma cidade equilibrada, tomando decisões sobre infraestrutura, serviços e sustentabilidade. Não há narrativa linear; a história emerge das consequências acumuladas das escolhas feitas ao longo da partida.", False, 11, None),
        66: ("Início da partida com atributos base.", False, 11, None),
        67: ("Carregamento e rolagem do dado físico.", False, 11, None),
        68: ("Movimentação do peão pelo tabuleiro.", False, 11, None),
        69: ("Decisão de compra, pagamento de aluguel ou aplicação de efeito especial.", False, 11, None),
        70: ("Atualização dos atributos, da renda recorrente e dos eventos de cidade.", False, 11, None),
        71: ("Encerramento após 12 rodadas com tela final e classificação da gestão.", False, 11, None),
        73: ("O mundo do jogo representa uma cidade estilizada organizada em torno de um tabuleiro central. O cenário de fundo mostra bairros temáticos, skyline urbano, áreas verdes e marcos visuais, enquanto o tabuleiro concentra as decisões estratégicas do jogador.", False, 11, None),
        75: ("Observação: além das propriedades acima, o jogo também inclui Usina Solar, Tratamento e Praça de Alimentação, cada uma com custos, renda e impactos próprios exibidos dinamicamente na interface de compra.", False, 11, None),
        77: ("O jogo é estruturado em um único tabuleiro jogável. O level design foca em clareza visual, leitura dos tipos de propriedade, percurso fechado com 20 casas e elementos de ambientação urbana ao fundo.", False, 11, None),
        79: ("O personagem principal é o gestor urbano representado por um peão físico no tabuleiro. Seu papel é tomar decisões sobre o desenvolvimento da cidade.", False, 11, None),
        80: ("As construções, propriedades e elementos urbanos funcionam como personagens sistêmicos, comunicando visualmente o impacto de cada escolha do jogador.", False, 11, None),
        82: ("Foram realizados testes iterativos no editor da Unity para validar rolagem do dado, movimentação do peão, funcionamento da câmera, compra de propriedades, exibição de construções, eventos de cidade, fluxo de HUD, menu, pausa e tela final.", False, 11, None),
        84: ("O projeto utiliza vetores (Vector3), interpolação, trigonometria e curvas de suavização para movimentação, câmera e animações. A pontuação final é calculada pela fórmula: Pontuação = Dinheiro + (Bem-estar × 10) - (Poluição × 8). O sistema também aplica impactos positivos e negativos por propriedade, eventos e renda recorrente.", False, 11, None),
        86: ("Unity Documentation.", False, 11, None),
        87: ("Nações Unidas. Objetivos de Desenvolvimento Sustentável.", False, 11, None),
        88: ("Assets utilizados no projeto: Daniel Riches - Board Game Essentials (dado D6 e peão TokenCup), Triangularity - Colorful City (prédios e árvores do cenário urbano), Palmov Island - Low Poly Houses Free Pack (casas, edifícios, banco, lixeira e roda-gigante), Studio Horizon - Simple Building Generic Free (escola e hospital), SolarPanel - GameReadyPrefab (painel solar) e Gwangju 3D Asset - SamhuiRestaurant (praça de alimentação).", False, 11, None),
        89: ("Referências de jogos de tabuleiro de gestão e planejamento urbano.", False, 11, None),
    }

    for idx, spec in updates.items():
        set_paragraph_text(doc.paragraphs[idx], spec[0], bold=spec[1], size=spec[2], align=spec[3])

    table = doc.tables[0]
    rows = [
        ["Tipo", "Compra", "Dinheiro/rodada", "Bem-estar", "Poluição"],
        ["Fábrica", "150", "+22", "-10", "+18"],
        ["Comércio", "150", "+18", "+5", "+8"],
        ["Residencial", "120", "+10", "+12", "+4"],
        ["Parque", "90", "-6", "+18", "-12"],
        ["Hospital", "170", "-6", "+22", "-4"],
        ["Escola", "170", "-6", "+16", "-4"],
    ]

    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            set_cell_text(table.rows[row_index].cells[col_index], value, bold_first=(row_index == 0))

    doc.save(GDD_PATH)
    doc.save(GDD_OUT_PATH)


def update_extension():
    doc = Document(EXT_SOURCE_PATH)
    tables = doc.tables

    set_cell_text(tables[0].cell(0, 0), "Eco City: jogo digital para conscientização sobre gestão urbana sustentável e tomada de decisão estratégica.")

    team_rows = [
        ("Arthur Lima De Luiz", "26028799"),
        ("Brian Walter", "26029357"),
        ("Leonardo Batista Franca", "26028512"),
    ]
    for idx, (name, ra) in enumerate(team_rows, start=1):
        set_cell_text(tables[1].cell(idx, 0), name)
        set_cell_text(tables[1].cell(idx, 1), ra)

    # Remove linhas extras caso existam nomes antigos
    for idx in range(len(team_rows) + 1, len(tables[1].rows)):
        set_cell_text(tables[1].cell(idx, 0), "")
        set_cell_text(tables[1].cell(idx, 1), "")

    set_cell_text(tables[2].cell(0, 0), "Professores orientadores: Victor Bruno Alexander Rosetti de Quiroz, Renata Muniz Do Nascimento, Adriano Felix Valente, Eduardo Savino Gomes e Luis Fernando dos Santos Pires.")
    set_cell_text(tables[3].cell(0, 0), "Ciência da Computação.")
    set_cell_text(tables[4].cell(0, 0), "Projeto Interdisciplinar.")
    set_cell_text(tables[4].cell(0, 1), "Linha de atuação centrada em tecnologia educacional, sustentabilidade urbana e conscientização cidadã.")
    set_cell_text(tables[5].cell(0, 0), "Atividade de Extensão não implementada na prática, estruturada como proposta de intervenção exequível com protótipo funcional desenvolvido e pronto para demonstração.")
    set_cell_text(tables[6].cell(0, 0), "ODS 11 - Cidades e Comunidades Sustentáveis; ODS 4 - Educação de Qualidade; ODS 9 - Indústria, Inovação e Infraestrutura; ODS 13 - Ação Contra a Mudança Global do Clima.")
    set_cell_text(tables[7].cell(0, 0), "Produto extensionista: jogo digital 3D Eco City, desenvolvido na Unity, com tabuleiro estratégico, dado físico, sistema de compra de propriedades, eventos urbanos, tela final classificatória, menu, HUD e documentação do projeto (GDD e relatório de extensão). Evidências previstas: capturas de tela, build executável, código-fonte e documentação. Assets utilizados no jogo: Board Game Essentials (dado e peão), Colorful City (skyline e vegetação), Low Poly Houses Free Pack (construções e detalhes urbanos), Simple Building Generic Free (escola e hospital), SolarPanel (usina solar) e SamhuiRestaurant (praça de alimentação).")
    set_cell_text(tables[8].cell(0, 0), "A intervenção pode ser aplicada em escolas de ensino médio, cursos técnicos, feiras acadêmicas, semanas de extensão e oficinas introdutórias sobre sustentabilidade e cidadania. O jogo foi pensado para um contexto educacional e demonstrativo, com execução viável em laboratórios de informática ou apresentações com projetor.")
    set_cell_text(tables[9].cell(0, 0), "O público-alvo é composto por estudantes de 15 a 25 anos, comunidade escolar, visitantes de eventos acadêmicos e pessoas interessadas em gestão urbana e sustentabilidade. Trata-se de um grupo com diferentes níveis de familiaridade com planejamento urbano, o que reforça a importância de uma abordagem lúdica, visual e acessível.")
    set_cell_text(tables[10].cell(0, 0), "O problema central identificado é a dificuldade de compreender, de forma prática e acessível, como decisões urbanas afetam simultaneamente economia, qualidade de vida e meio ambiente. Em geral, esses temas são discutidos de forma abstrata, sem ferramentas de experimentação. O projeto responde a esse problema por meio de um jogo que transforma esses conceitos em escolhas concretas e consequências imediatas.")
    set_cell_text(tables[11].cell(0, 0), "Hipótese de intervenção: um jogo digital educativo, com regras simples e feedback visual claro, pode ampliar o entendimento do público sobre gestão urbana sustentável. Ao experimentar decisões de compra, renda, poluição, bem-estar e eventos da cidade, o participante tende a perceber melhor os impactos de políticas urbanas e a importância do equilíbrio entre desenvolvimento econômico e responsabilidade socioambiental.")
    set_cell_text(tables[12].cell(0, 0), "Eco City é uma proposta de projeto de extensão baseada no desenvolvimento de um jogo digital 3D de estratégia, criado na Unity, para sensibilizar estudantes e a comunidade sobre sustentabilidade urbana. O jogo coloca o participante no papel de gestor de uma cidade, que precisa tomar decisões ao longo de 12 rodadas em um tabuleiro com 20 casas. Cada propriedade adquirida gera efeitos em dinheiro, bem-estar e poluição, além de eventos especiais e múltiplos finais. A proposta pode ser aplicada em oficinas, mostras acadêmicas e ações educativas, funcionando como ferramenta lúdica de discussão sobre planejamento urbano, equilíbrio ambiental e responsabilidade pública.")
    set_cell_text(tables[13].cell(0, 0), "A extensão universitária busca aproximar o conhecimento acadêmico das demandas sociais, criando intervenções com impacto formativo e comunitário. Nesse contexto, Eco City propõe um jogo digital como instrumento educativo voltado à compreensão de desafios urbanos contemporâneos. O projeto dialoga diretamente com a ODS 11, ao abordar cidades sustentáveis, e também com a ODS 4, ao favorecer aprendizagem significativa, com a ODS 9, por estimular inovação tecnológica, e com a ODS 13, ao discutir impactos ambientais. A fundamentação do projeto parte da ideia de que jogos sérios e experiências interativas podem ampliar o engajamento dos participantes, facilitando a compreensão de sistemas complexos e relações de causa e efeito. Em Eco City, o jogador observa na prática como escolhas econômicas, sociais e ambientais precisam ser equilibradas para produzir uma cidade saudável e funcional.")
    set_cell_text(tables[14].cell(0, 0), "• Desenvolver um jogo digital 3D com foco em gestão urbana sustentável.\n• Simular decisões sobre propriedades, renda, poluição e bem-estar em um tabuleiro estratégico.\n• Estimular o pensamento crítico sobre políticas públicas urbanas e sustentabilidade.\n• Produzir um artefato extensionista aplicável em contextos educativos e acadêmicos.\n• Promover discussão sobre equilíbrio entre desenvolvimento econômico, qualidade de vida e preservação ambiental.")
    set_cell_text(tables[15].cell(0, 0), "A metodologia combina pesquisa, design de jogo, implementação técnica e validação prática. Inicialmente, a equipe levantou referências sobre sustentabilidade urbana, jogos educativos e mecânicas de tabuleiro. Em seguida, estruturou o fluxo principal do jogo, definindo atributos, rodada, progressão, eventos especiais e finais classificatórios. A implementação foi realizada na Unity com C#, incluindo dado físico com carregamento, movimentação do peão, compra de propriedades, construções visuais, renda recorrente, HUD, menus, câmera dinâmica e tela final. Na composição visual, foram integrados assets de terceiros para apoiar o protótipo, como peão e dado, prédios low poly, árvores, escola, hospital, painel solar e restaurante temático. Para a etapa extensionista, a proposta é utilizar o jogo em demonstrações, oficinas e apresentações, permitindo que os participantes joguem, discutam as consequências das decisões e relacionem a experiência com problemas reais das cidades. A coleta de percepção pode ocorrer por observação, conversa guiada e questionários simples após a atividade.")
    set_cell_text(tables[16].cell(0, 0), "Resultados esperados: maior compreensão do público sobre os efeitos interligados de decisões urbanas; fortalecimento do debate sobre sustentabilidade, cidadania e planejamento urbano; produção de um jogo educacional funcional e apresentável; e geração de material acadêmico que comprove o desenvolvimento do projeto. Como resultado parcial já alcançado, o jogo possui menu inicial, HUD, dado físico, movimentação do peão, compra de propriedades, eventos de cidade, cenário urbano, sistema de pontuação e tela final classificatória.")
    set_cell_text(tables[17].cell(0, 0), "Conclui-se que Eco City constitui uma proposta consistente de extensão universitária por transformar conceitos de sustentabilidade urbana em uma experiência prática e interativa. O projeto responde ao problema central ao criar uma ferramenta acessível para discutir planejamento urbano, impacto ambiental e tomada de decisão estratégica. Além do valor técnico no desenvolvimento do software, o jogo apresenta potencial educativo e extensionista para futuras ações em escolas, feiras e eventos acadêmicos.")
    set_cell_text(tables[18].cell(0, 0), "NAÇÕES UNIDAS. Objetivos de Desenvolvimento Sustentável. Disponível em: https://brasil.un.org/pt-br/sdgs.\nUNITY TECHNOLOGIES. Unity Manual e Scripting API. Disponível em: https://docs.unity3d.com/.\nSALEN, Katie; ZIMMERMAN, Eric. Rules of Play: Game Design Fundamentals.\nReferências complementares sobre jogos educativos, sustentabilidade urbana e aprendizagem baseada em simulação.")
    set_cell_text(tables[19].cell(0, 0), "Produto extensionista principal: game digital educativo Eco City.\nProdutos complementares: GDD atualizado, documentação do projeto, imagens de tela, possível build executável, apresentação acadêmica e roteiro para demonstração em oficinas e mostras.")
    set_cell_text(tables[20].cell(0, 0), "Fontes utilizadas\n• Objetivos de Desenvolvimento Sustentável\n• Documentação oficial da Unity\n• Referências de game design e jogos sérios")
    set_cell_text(tables[20].cell(0, 1), "https://brasil.un.org/pt-br/sdgs\nhttps://docs.unity3d.com/\nhttps://unity.com/")
    set_cell_text(tables[21].cell(0, 1), "https://www.fecap.br/")
    set_cell_text(tables[21].cell(1, 1), "https://www.fecap.br/")

    doc.save(EXT_PATH)
    doc.save(EXT_OUT_PATH)


if __name__ == "__main__":
    update_gdd()
    update_extension()
    print("Documentos atualizados com UTF-8.")
